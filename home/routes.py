import math
import overpy
import requests
import io
import logging
import os
from datetime import datetime
from io import BytesIO
from urllib.parse import unquote
from functools import wraps
import googlemaps
import numpy as np
import tensorflow as tf
from PIL import Image
from docx import Document
from flask import jsonify, render_template, url_for, redirect, request, flash, send_file, make_response, Response
from flask_login import current_user, login_required, login_user, logout_user
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from werkzeug.utils import secure_filename
from tensorflow.keras.losses import SparseCategoricalCrossentropy
from home import app, db, login, cache
from home.forms import LoginForm, RegistrationForm, UploadForm, EditProfileForm, CropRecommendationForm
from home.models import User, SQLAlchemyError, UploadedImage, Prediction


@app.before_request
def before_request():
    if current_user.is_authenticated:
        current_user.last_seen = datetime.utcnow()
        db.session.commit()


GOOGLE_MAPS_API_KEY = os.getenv('GOOGLE_MAPS_API_KEY')
HERE_API_KEY = os.getenv('HERE_API_KEY') #HERE API KEY, AN ALTERNATIVE FOR THE GEOCODING OF THE NEAREST AGROVETS

# Initialize API clients
gmaps = googlemaps.Client(key=GOOGLE_MAPS_API_KEY) if GOOGLE_MAPS_API_KEY else None

# Check for API key loading
if not GOOGLE_MAPS_API_KEY:
    flash("Google Maps API key is missing. Please check your .env file.")


def predict_plant_class(image_bytes):
    try:
        model_path = 'model.h5'
        if os.path.exists(model_path):
            model = tf.keras.models.load_model(model_path, custom_objects={
                'SparseCategoricalCrossentropy': SparseCategoricalCrossentropy()  # Custom loss function
            })
        else:
            return "Failed to load the plant class model. Please check the model path."

        image = tf.image.decode_jpeg(image_bytes.getvalue(), channels=3)
        image = tf.image.resize(image, [128, 128])
        image = tf.expand_dims(image, axis=0)
        image = tf.cast(image, tf.float32)

        predictions = model.predict(image)
        class_names = ['Cashew', 'Cassava', 'Maize', 'Tomato']
        predicted_class = class_names[np.argmax(predictions)]

        return predicted_class
    except AttributeError:
        return "Error: Please provide a valid image."
    except Exception as e:
        error_message = f"An error occurred during plant class prediction: {e}"
        logging.error(error_message)
        return error_message


# Placeholder for model prediction function for disease, pest, or healthy status
def predict_disease_or_pest_or_healthy(image_bytes, plant_class):
    """Placeholder for model prediction function for disease, pest, or healthy status."""
    try:
        # Load the model for predicting disease, pest, or healthy status
        model_path = 'best_model22.h5'
        if os.path.exists(model_path):
            model = tf.keras.models.load_model(model_path)
        else:
            return "Failed to load the disease, pest, or healthy model. Please check the model path."

        # Preprocess the image bytes for disease, pest, or healthy status prediction
        image = tf.image.decode_jpeg(image_bytes.getvalue(), channels=3)
        image = tf.image.resize(image, [224, 224])  # Resize the image
        image = tf.expand_dims(image, axis=0)  # Add batch dimension
        image = tf.cast(image, tf.float32)  # Convert image to float32

        # Placeholder logic to determine disease, pest, or healthy status
        predictions = model.predict(image)  # Make predictions based on the preprocessed input

        class_names = [
            'Cashew anthracnose', 'Cashew healthy', 'Cashew leaf miner', 'Cashew gumosis', 'Cashew red rust',  # Cashew
            'Cassava bacterial blight', 'Cassava brown spot', 'Cassava green mite', 'Cassava healthy', 'Cassava mosaic',
            # Cassava
            'Maize fall armyworm', 'Maize grasshopper', 'Maize healthy', 'Maize leaf beetle', 'Maize leaf blight',
            'Maize leaf spot', 'Maize streak virus',  # Maize
            'Tomato healthy', 'Tomato leaf blight', 'Tomato leaf curl', 'Tomato septoria leaf spot',
            'Tomato verticulium wilt'  # Tomato
        ]
        predicted_class = class_names[np.argmax(predictions)]  # Decode the predicted output

        return predicted_class
    except AttributeError as e:
        # Handle the specific error when image_bytes is None
        error_message = "Error: Please provide a valid image."
        return error_message
    except Exception as e:
        error_message = f"An error occurred during disease, pest, or healthy status prediction: {e}"
        logging.error(error_message)  # Log the error
        return error_message


image_form = ['jpeg', 'jpg', 'png', 'webp']


def load_and_preprocess_image(image_path):
    """Load and preprocess the image for prediction."""
    try:
        # Open the image file
        img = Image.open(image_path)
        # Resize the image
        img = img.resize((300, 300))
        # Normalize the image and convert it to a numpy array
        img_array = np.array(img) / 255.0
        # Reshape the image array to match the input shape of the model
        img_array = img_array.reshape((img_array.shape[0], img_array.shape[1], 3))

        # Convert the image array to bytes and store it in a BytesIO object
        image_bytes = io.BytesIO()
        Image.fromarray((img_array * 255).astype(np.uint8)).save(image_bytes, format='JPEG')

        return image_bytes
    except Exception as e:
        return None


# Function to read the disease explanation from the file
def read_disease_explanation(predicted_disease):
    try:
        # Construct the filename based on the predicted disease
        filename = predicted_disease.lower().replace(' ', '_') + '.txt'
        explanation_file_path = os.path.join(os.path.dirname(__file__), 'explained', filename)

        # Check if the file exists
        if os.path.exists(explanation_file_path):
            with open(explanation_file_path, 'r') as f:
                explanation = f.read()
            return explanation
        else:
            # Return a default message if the explanation file is not found
            return "Description not available for this disease."
    except Exception as e:
        return f"An error occurred while reading the disease explanation: {e}"


# Function to get the first 50 words of the explanation
def get_first_50_words(explanation):
    words = explanation.split()[:50]
    return ' '.join(words)


def allowed_file(filename):
    return '.' in filename and \
        filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def health_check_route(route_func):
    @wraps(route_func)
    def wrapper(*args, **kwargs):
        try:
            # Execute the route function
            response = route_func(*args, **kwargs)
            
            # Check if the response is valid
            if response.status_code == 200:
                return jsonify({'status': 'ok'}), 200
            else:
                return jsonify({'status': 'error', 'message': 'Route not healthy'}), 500
        except Exception as e:
            return jsonify({'status': 'error', 'message': str(e)}), 500
    return wrapper

from flask import jsonify, flash, render_template, redirect, url_for

@app.route('/predict', methods=['GET', 'POST'])
@login_required
def predict():
    form = UploadForm()

    try:
        if form.validate_on_submit():
            if 'image' not in request.files:
                flash('No image uploaded', 'error')
                return redirect(url_for('predict'))

            file = form.image.data

            if file.filename == '':
                flash('No selected image', 'error')
                return redirect(url_for('predict'))

            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)

                # Save the uploaded image
                uploaded_image = UploadedImage(filename=filename)
                db.session.add(uploaded_image)
                db.session.commit()

                # Load and preprocess the image
                image = load_and_preprocess_image(filepath)
                plant_class = predict_plant_class(image)

                if not plant_class:
                    flash('Failed to predict plant class', 'error')
                    return redirect(url_for('predict'))

                # Predict disease or pest or healthy status
                disease_or_pest_or_healthy = predict_disease_or_pest_or_healthy(image, plant_class)

                if not disease_or_pest_or_healthy:
                    flash('Failed to predict disease, pest, or healthy status', 'error')
                    return redirect(url_for('predict'))

                # Read disease explanation
                disease_explanation = read_disease_explanation(disease_or_pest_or_healthy)
                first_50_words = get_first_50_words(disease_explanation)

                # Create a prediction instance
                prediction = Prediction(
                    plant_class=plant_class,
                    disease_status=disease_or_pest_or_healthy,
                    first_50_words=first_50_words,
                    disease_explanation=disease_explanation,
                    author=current_user,
                    image=uploaded_image
                )
                db.session.add(prediction)
                db.session.commit()

                flash('Prediction successful!', 'success')

                # Rendering template with additional variables
                return render_template('predict.html', form=form, plant_class=plant_class,
                                       disease_status=disease_or_pest_or_healthy,
                                       first_50_words=first_50_words,
                                       disease_explanation=disease_explanation)

            else:
                flash('Invalid file type. Allowed file types are: png, jpg, jpeg, gif', 'error')
                return redirect(url_for('predict'))

        elif request.method == 'GET':
            return render_template('predict.html', form=form)

    except Exception as e:
        error_message = f'An unexpected error occurred: {str(e)}'
        logging.error(error_message)
        flash('An unexpected error occurred. Please try again later.', 'error')
        return redirect(url_for('predict'))

    return render_template('predict.html', form=form)


def handle_error(message, status_code, is_json=False):
    if is_json:
        return jsonify({'error': message}), status_code
    else:
        flash(message, 'danger')
        return redirect(url_for('predict'))

def handle_error(message, status_code, is_json=False):
    if is_json:
        response = jsonify({'status': 'error', 'message': message})
    else:
        response = make_response(message, status_code)
    response.status_code = status_code
    return response

@app.route('/download-explanation/<predicted_disease>', methods=['GET'])
@login_required
def download_explanation(predicted_disease):
    try:
        # Remove newline characters from the predicted disease name
        predicted_disease = unquote(predicted_disease)

        # Construct the filename based on the predicted disease
        filename = predicted_disease.lower().replace(' ', '_') + '.txt'
        explanation_file_path = os.path.join(os.path.dirname(__file__), 'explained', filename)

        # Check if the file exists
        if os.path.exists(explanation_file_path):
            return send_file(explanation_file_path, as_attachment=True,
                             download_name=f"{predicted_disease.replace(' ', '_')}_explanation.txt")
        else:
            # Return a default message if the explanation file is not found
            error_message = f'Description not available for this disease.'
            logging.error(error_message)
            flash('Description not available for this disease.', 'danger')
            return redirect(url_for('predict'))
    except Exception as e:
        error_message = f'An unexpected error occurred while downloading the explanation file: {str(e)}'
        logging.error(error_message)
        flash('An unexpected error occurred while downloading the explanation file. Please try again later.', 'danger')
        return redirect(url_for('predict'))


# Function to generate a Word document from the disease explanation
def generate_word_document(explanation):
    document = Document()
    document.add_heading('Disease Explanation', level=1)
    document.add_paragraph(explanation)
    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


def generate_pdf_document(explanation):
    styles = getSampleStyleSheet()
    style_heading = styles['Heading1']
    style_normal = styles['Normal']

    story = []
    story.append(Paragraph("Disease Explanation", style_heading))
    story.append(Spacer(1, 12))

    # Split the explanation into sections (if applicable) and add them as paragraphs
    sections = explanation.split('\n\n')
    for section in sections:
        story.append(Paragraph(section.strip(), style_normal))
        story.append(Spacer(1, 6))  # Add some spacing between sections

    pdf_buffer = BytesIO()
    doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
    doc.build(story)
    pdf_buffer.seek(0)
    return pdf_buffer


# Route function to handle the download request for the Word document
@app.route('/download-explanation-word/<predicted_disease>')
def download_explanation_word(predicted_disease):
    disease_explanation = read_disease_explanation(predicted_disease)
    word_output = generate_word_document(disease_explanation)
    return send_file(word_output, as_attachment=True,
                     download_name=f"{predicted_disease.replace(' ', '_')}_explanation.docx")


# Route function to handle the download request for the PDF document
@app.route('/download-explanation-pdf/<predicted_disease>')
def download_explanation_pdf(predicted_disease):
    disease_explanation = read_disease_explanation(predicted_disease)
    pdf_output = generate_pdf_document(disease_explanation)
    return send_file(pdf_output, as_attachment=True,
                     download_name=f"{predicted_disease.replace(' ', '_')}_explanation.pdf")


@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('predict'))
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()
        if user is None or not user.check_password(form.password.data):
            flash('Invalid username or password')
            return redirect(url_for('login'))
        login_user(user, remember=form.remember_me.data)
        flash(f'Welcome {user.username}')
        return redirect(url_for('predict'))
    return render_template(
        'login.html',
        title='Login',
        form=form)


@app.route('/logout')
def logout():
    logout_user()
    return redirect(url_for('login'))


@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('predict'))
    form = RegistrationForm()
    if form.validate_on_submit():
        user = User(
            username=form.username.data,
            email=form.email.data,
        )
        try:
            user.set_password(form.password.data)
            db.session.add(user)
            db.session.commit()
            flash('Congratulations! Login to continue')
            return redirect(url_for('login'))
        except SQLAlchemyError as e:
            db.session.rollback()
            logging.error(f'Error registering: {str(e)}')
            flash('An error occurred while saving the user. Please try again.', 'danger')
    return render_template(
        'register.html',
        title='Register',
        form=form
    )


# Cache the results of conversation queries
@cache.cached(timeout=300)  # Cache for 5 minutes (adjust as needed)
def get_history(user_id):
    user = User.query.get(user_id)
    return user.predictions.order_by(Prediction.timestamp.desc()).all()


# Invalidate the cache whenever new conversations are added or updated
def invalidate_conversation_cache(user_id):
    cache.delete_memoized(get_history, user_id)


@app.route('/')
def home():
    # Dummy data for carousel items (replace with your actual data)
    carousel_items = [
        {
            'image_url': url_for('static', filename='/images/coffee_home.jpg'),
            'heading': 'Welcome to Coffee AgroHealth',
            'description': 'Coffee Agriculture utility website to help farmers by providing harvesting solutions using artificial intelligence.'
        },
        # Add more carousel items as needed
    ]
    return render_template('index.html', carousel_items=carousel_items)


# Implement lazy loading combined with pagination
@app.route('/user/<username>')
@login_required
def user(username):
    user = User.query.filter_by(username=username).first_or_404()
    page = request.args.get('page', 1, type=int)
    prediction_history_pagination = user.predictions.order_by(Prediction.timestamp.desc()).paginate(
        page=page, per_page=app.config['HISTORY_PER_PAGE'], error_out=False)

    # Fetch conversations for the current page
    predictions = prediction_history_pagination.items
    print(predictions)
    next_url = url_for('user', username=username, page=prediction_history_pagination.next_num) \
        if prediction_history_pagination.has_next else None
    prev_url = url_for('predict', username=username, page=prediction_history_pagination.prev_num) \
        if prediction_history_pagination.has_prev else None

    return render_template(
        'user.html',
        title='Prediction History',
        username=username,
        user=user,
        next_url=next_url,
        prev_url=prev_url,
        predictions=predictions
    )


@app.route('/edit-profile', methods=['GET', 'POST'])
@login_required
def edit_profile():
    form = EditProfileForm(current_user.username)
    if form.validate_on_submit():
        current_user.username = form.username.data
        current_user.about_me = form.about_me.data
        db.session.commit()
        flash('Your changes have been saved.')
        return redirect(url_for('user', username=current_user.username))
    elif request.method == 'GET':
        form.username.data = current_user.username
        form.about_me.data = current_user.about_me
    return render_template(
        'edit_profile.html',
        title='Edit Profile',
        form=form
    )


@app.route('/user/<username>/popup')
@login_required
def user_popup(username):
    user = User.query.filter_by(username=username).first_or_404()
    return render_template(
        'user_popup.html',
        user=user)


@app.route('/about_us')
def about_us():
    return render_template('about_us.html')


def display_image_grid(image_paths, num_cols=4):
    num_images = len(image_paths)
    num_rows = -(-num_images // num_cols)  # Ceiling division to calculate the number of rows

    grid = []
    for i in range(num_rows):
        row = []
        for j in range(num_cols):
            index = i * num_cols + j
            if index < num_images:
                row.append(image_paths[index])
            else:
                break
        grid.append(row)
    return grid


@app.route('/grid_images')
def grid_images():
    image_paths = [
        "static/images/image1.jpg",
        "static/images/image2.webp",
        "static/images/image3.jpg",
        "static/images/image4.jpg",
        "static/images/image5.jpg",
        "static/images/image9.jpeg",
        "static/images/image7.jpeg",
        "static/images/image8.jpeg",
        "static/images/image9.jpg",
        "static/images/image10.jpg",
        "static/images/image11.jpg",
        "static/images/image12.jpg",
    ]
    image_grid = display_image_grid(image_paths)
    return render_template('image_grid.html', image_grid=image_grid)

# Function to calculate distance between two points using Haversine formula
def haversine(lat1, lon1, lat2, lon2):
    R = 6371  # Radius of the Earth in kilometers
    dLat = math.radians(lat2 - lat1)
    dLon = math.radians(lon2 - lon1)
    a = math.sin(dLat / 2) * math.sin(dLat / 2) + math.cos(math.radians(lat1)) * math.cos(
        math.radians(lat2)) * math.sin(dLon / 2) * math.sin(dLon / 2)
    c = 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))
    distance = R * c
    return distance

# Geocode location using Google Geocoding API
GEOCODING_API_KEY = os.getenv('GEOCODING_API_KEY')
def geocode_location(location):
    url = f"https://maps.googleapis.com/maps/api/geocode/json?address={location}&key={GEOCODING_API_KEY}"
    response = requests.get(url)
    data = response.json()
    if data['status'] == 'OK':
        # Extract latitude and longitude from the geocoding response
        lat = data['results'][0]['geometry']['location']['lat']
        lon = data['results'][0]['geometry']['location']['lng']
        return lat, lon
    else:
        return None, None

# Get closest agrovet or agrocenter to the user's location
def get_closest_agrovet(latitude, longitude):
    api = overpy.Overpass()
    query = """
        node["shop"="agrovet"](around:5000,{lat},{lon});
        out;
    """.format(lat=latitude, lon=longitude)
    result = api.query(query)

    user_location = (latitude, longitude)  # Storing the user's location
    closest_agrovet = None
    min_distance = float('inf')

    for node in result.nodes:
        agrovet_location = (node.lat, node.lon)
        distance = haversine(latitude, longitude, node.lat, node.lon)
        if distance < min_distance:
            min_distance = distance
            closest_agrovet = agrovet_location

    return closest_agrovet, user_location  # Returning the user's location
# USING HERE API FOR GEO-LOCATION OF THE AGROVETS

# @app.route('/argo_location', methods=["GET", "POST"])
# def get_nearby_agrovet_services():
#     endpoint = f'https://discover.search.hereapi.com/v1/discover'
#     if request.method == "POST":
#         location = request.form['location']
#         latitude, longitude = geocode_location(location)
#     # Define query parameters
#     params = {
#         'q': 'agrovet',
#         'at': location,  # Location coordinates
#         'limit': 6,
#         'apikey': HERE_API_KEY
#     }
    
#     # Make the request to HERE API
#     response = requests.get(endpoint, params=params)
#     data = response.json()
    
#     # Extract relevant information
#     if 'items' in data:
#         services = []
#         for item in data['items']:
#             name = item['title']
#             address = item.get('address', {}).get('label', 'No address provided')
#             services.append(f"{name} located at {address}")
#         return return render_template('results.html', closest_agrovet=services, lat=latitude, lon=longitude,
#                                        location=location, user_location=user_location)
#     else:
#         return render_template('not_found.html', lat=latitude, lon=longitude, user_location=user_location)

@app.route('/argo_location', methods=["GET", "POST"])
@login_required
def argo_location():
    if request.method == "POST":
        # Your existing code for handling POST requests
        location = request.form["location"]
        latitude, longitude = geocode_location(location)

        if latitude is not None and longitude is not None:
            closest_agrovet, user_location = get_closest_agrovet(latitude, longitude)
            if closest_agrovet is not None:
                # Render template with closest agrovet
                return render_template('results.html', closest_agrovet=closest_agrovet, lat=latitude, lon=longitude,
                                       location=location, user_location=user_location)
            else:
                # Render template with user's location and not found message
                return render_template('not_found.html', lat=latitude, lon=longitude, user_location=user_location)
        else:
            return "Location not found. Please enter a valid location.", 400
    elif request.method == "GET":
        # Your code for handling GET requests
        return render_template('input.html')  # or any other appropriate response for GET requests
    else:
        # Handling wrong method error for any other methods
        return "Method not allowed. Please use the POST method for this endpoint.", 405

# Health check route
@app.route('/health')
def health_check():
    return jsonify({'status': 'ok', 'time': datetime.utcnow()})
<<<<<<< HEAD
=======


# object detection section
import cv2
from ultralytics import YOLO
import threading
import time
import logging
import io
from PIL import Image
import time
import psutil
import platform

# Load the YOLOv8 model using ultralytics
model = YOLO('./home//model/best.pt')

# Global variable to control detection and store the last frame
detection_enabled = True
last_detected_frame = None
frame_lock = threading.Lock()

class VideoCaptureThread:
    def __init__(self, src=0):
        self.cap = cv2.VideoCapture(src)
        self.frame = None
        self.ret = False
        self.thread = threading.Thread(target=self.update, args=())
        self.thread.daemon = True
        self.thread.start()

    def update(self):
        while True:
            if self.cap.isOpened():
                self.ret, self.frame = self.cap.read()
            else:
                break

    def read(self):
        return self.ret, self.frame

    def release(self):
        if self.cap.isOpened():
            self.cap.release()

def gen_frames(skip_frames=10, target_size=(480, 320), delay=0.07):
    cap = VideoCaptureThread(0)  # Threaded video capture
    frame_counter = 0

    try:
        while True:
            if not detection_enabled:
                # Serve the last detected frame if available
                with frame_lock:
                    if last_detected_frame is not None:
                        ret, buffer = cv2.imencode('.jpg', last_detected_frame)
                        if not ret:
                            logging.error("Failed to encode last detected frame.")
                            continue
                        frame = buffer.tobytes()
                        yield (b'--frame\r\n'
                               b'Content-Type: image/jpeg\r\n\r\n' + frame + b'\r\n')
                time.sleep(delay)
                continue

            ret, frame = cap.read()  # Read frame from the camera
            if not ret:
                logging.error("Failed to capture frame.")
                continue

            # Resize the frame to reduce processing time
            frame_resized = cv2.resize(frame, target_size)

            # Skip frames
            if frame_counter % skip_frames == 0:
                # Perform inference
                results = model(frame_resized)
                annotated_frame = results[0].plot()  # Get annotated frame

                # Store the last detected frame
                with frame_lock:
                    last_detected_frame = annotated_frame.copy()

                # Encode the frame as JPEG
                ret, buffer = cv2.imencode('.jpg', annotated_frame)
                if not ret:
                    logging.error("Failed to encode frame.")
                    continue
                frame = buffer.tobytes()

                yield (b'--frame\r\n'
                       b'Content-Type: image/jpeg\r\n\r\n' + frame + b'\r\n')

                # Introduce a delay for better rendering
                time.sleep(delay)

            frame_counter += 1
    finally:
        # Release camera resource when done
        cap.release()
        cv2.destroyAllWindows()

def update(self):
    while True:
        if self.cap.isOpened():
            self.ret, self.frame = self.cap.read()
            if not self.ret:
                logging.error("Failed to read frame.")
        else:
            logging.error("Camera is not opened.")
            break

@app.route('/coffee_detector')
@login_required
def index():
    return render_template('object_detection.html')

@app.route('/video_feed')
@login_required
def video_feed():
    return Response(gen_frames(), mimetype='multipart/x-mixed-replace; boundary=frame')

@app.route('/toggle_detection', methods=['POST'])
@login_required
def toggle_detection():
    global detection_enabled
    detection_enabled = not detection_enabled
    status = 'enabled' if detection_enabled else 'disabled'
    logging.error(f"Detection toggled {status}")
    logging.info(f"Detection toggled {status}.")
    return jsonify({"status": status})


@app.route('/capture_snapshot', methods=['POST'])
@login_required
def capture_snapshot():
    global last_detected_frame
    if last_detected_frame is None:
        return jsonify({"status": "error", "message": "No frame available to capture."})

    # Convert the frame to a PIL image and save it to a bytes buffer
    img = Image.fromarray(last_detected_frame)
    buffer = io.BytesIO()
    img.save(buffer, format="JPEG")
    buffer.seek(0)

    # Return the image as a file
    return Response(buffer, mimetype='image/jpeg', headers={"Content-Disposition": "attachment;filename=snapshot.jpg"})


def get_performance_metrics():
    # Get system performance metrics
    cpu_usage = psutil.cpu_percent()
    memory_info = psutil.virtual_memory()
    return {
        "cpu_usage": cpu_usage,
        "memory_usage": memory_info.percent,
        "os": platform.system() + " " + platform.version()
    }

@app.route('/performance_metrics')
@login_required
def performance_metrics():
    metrics = get_performance_metrics()
    return jsonify(metrics)
>>>>>>> c3943ad (first commit)
